import os
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import io


async def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(file_path)
        text = ""
        page_count = len(reader.pages)
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip(), page_count
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


async def extract_text_from_docx(file_path: str) -> tuple[str, int]:
    """Extract text from a DOCX file."""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        page_count = len(doc.paragraphs)  # Approximate
        
        return text.strip(), page_count
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


async def extract_text_from_file(file_path: str, file_type: str) -> tuple[str, int]:
    """Extract text from a file based on its type."""
    if file_type == "application/pdf":
        return await extract_text_from_pdf(file_path)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return await extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Only PDF and DOCX (Office 2007+) files are supported.")


async def save_uploaded_file(file_content: bytes, filename: str, upload_dir: str) -> str:
    """Save an uploaded file to disk."""
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, filename)
    
    with open(file_path, "wb") as f:
        f.write(file_content)
    
    return file_path


async def delete_file(file_path: str) -> bool:
    """Delete a file from disk."""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            return True
        return False
    except Exception as e:
        print(f"Error deleting file: {e}")
        return False
